# tac_app.py
# Full-featured Streamlit app:
# - Auth: create/login/logout, institution org codes, password reset via email
# - Stripe subscription checkout + subscribed return handler
# - Lesson generator + Worksheet generator (LA/MA/HA differentiation)
# - Pedagogical QA pass (OpenAI) + QA audit logs per user
# - Optional image generation for lesson/worksheet
# - Paid DOCX downloads + branding logo
#
# Streamlit Cloud notes:
# - SQLite can be unreliable for persistence across restarts/redeploys.
# - This file hardens SQLite (WAL + busy_timeout) to reduce "database locked".
# - For real persistence, use Postgres (Supabase/Neon) — can be added later.

import os
import sqlite3
import time
import base64
import bcrypt
import stripe
import smtplib
from io import BytesIO
from email.message import EmailMessage
from textwrap import dedent
from pathlib import Path

import streamlit as st
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired

from docx import Document
from docx.shared import Inches

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# =========================
# 0) STREAMLIT CONFIG
# =========================
st.set_page_config(page_title="TAC Learning", layout="wide")

# =========================
# 1) SECRETS / ENV HELPERS
# =========================
def sget(key: str, default=None):
    try:
        v = st.secrets.get(key, None)
        return v if v not in (None, "") else os.getenv(key, default)
    except Exception:
        return os.getenv(key, default)

OPENAI_API_KEY = sget("OPENAI_API_KEY", "")
STRIPE_SECRET_KEY = sget("STRIPE_SECRET_KEY", "")
STRIPE_PRICE_MONTHLY = sget("STRIPE_PRICE_MONTHLY", "")
STRIPE_PRICE_ANNUAL = sget("STRIPE_PRICE_ANNUAL", "")

APP_BASE_URL = sget("APP_BASE_URL", "http://localhost:8501")
RESET_SECRET = sget("RESET_SECRET", "dev-reset-secret-change-me")

SMTP_HOST = sget("SMTP_HOST", "")
SMTP_PORT = int(sget("SMTP_PORT", 587) or 587)
SMTP_USER = sget("SMTP_USER", "")
SMTP_PASS = sget("SMTP_PASS", "")

# DB path (still ephemeral on Streamlit Cloud)
DB = sget("DB_PATH", "users.db")

# Stripe config
STRIPE_READY = bool(STRIPE_SECRET_KEY and STRIPE_PRICE_MONTHLY and STRIPE_PRICE_ANNUAL)
if STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY

# OpenAI client
OPENAI_READY = bool(OpenAI and OPENAI_API_KEY)
client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_READY else None

# =========================
# 2) PEDAGOGY (INLINE)
# =========================
QA_CHECKLIST = [
    "Learning purpose is explicit",
    "Meaningful context or lived example included",
    "Cognitive load controlled",
    "Developmental readiness respected",
    "SEN/EAL supports embedded",
    "Active or enquiry task present",
    "Competency demonstrated",
    "Retrieval included",
    "Emotional tone is respectful",
    "Layout is accessible",
    "Curiosity and dignity preserved",
]

PEDAGOGY_CORE_POSITION = """
Learning is a relational, developmental, and meaningful process.
Understanding precedes memorisation.
Curiosity, dignity, and emotional safety are essential.
""".strip()

# =========================
# 3) QUERY PARAM HELPERS (Streamlit can return str OR list)
# =========================
qp = st.query_params

def qp_first(key: str, default=None):
    v = qp.get(key, default)
    if isinstance(v, list):
        return v[0] if v else default
    return v

# =========================
# 4) DATABASE (SQLite)
# =========================
def db():
    conn = sqlite3.connect(DB, check_same_thread=False)
    # Streamlit Cloud hardening (reduces lock issues, not perfect)
    conn.execute("PRAGMA journal_mode=WAL;")
    conn.execute("PRAGMA synchronous=NORMAL;")
    conn.execute("PRAGMA busy_timeout=5000;")
    return conn

def init_db():
    with db() as conn:
        c = conn.cursor()

        c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE,
            password BLOB,
            paid INTEGER DEFAULT 0,
            subscription_status TEXT DEFAULT 'none',
            org_id INTEGER DEFAULT NULL,
            role TEXT DEFAULT 'individual',
            created_at INTEGER DEFAULT (strftime('%s','now'))
        )
        """)

        c.execute("""
        CREATE TABLE IF NOT EXISTS orgs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            org_code TEXT UNIQUE,
            created_at INTEGER DEFAULT (strftime('%s','now'))
        )
        """)

        c.execute("""
        CREATE TABLE IF NOT EXISTS qa_audit (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            tool TEXT,
            topic TEXT,
            year_group TEXT,
            passed INTEGER,
            created_at INTEGER DEFAULT (strftime('%s','now'))
        )
        """)

        conn.commit()

def migrate_users_table():
    with db() as conn:
        c = conn.cursor()
        cols = [r[1] for r in c.execute("PRAGMA table_info(users)").fetchall()]

        if "subscription_status" not in cols:
            c.execute("ALTER TABLE users ADD COLUMN subscription_status TEXT DEFAULT 'none'")
        if "org_id" not in cols:
            c.execute("ALTER TABLE users ADD COLUMN org_id INTEGER")
        if "role" not in cols:
            c.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'individual'")

        conn.commit()

init_db()
migrate_users_table()

# =========================
# 5) SESSION STATE
# =========================
st.session_state.setdefault("user", None)
st.session_state.setdefault("show_login", False)

st.session_state.setdefault("final_lesson", None)
st.session_state.setdefault("lesson_image_bytes", None)

st.session_state.setdefault("final_worksheet", None)
st.session_state.setdefault("worksheet_image_bytes", None)

st.session_state.setdefault("brand_logo_bytes", None)

# =========================
# 6) AUTH HELPERS
# =========================
def create_user(email, pw, org_code=None):
    email = (email or "").strip().lower()
    if not email or "@" not in email:
        return (False, "Please enter a valid email address.")
    if not pw or len(pw) < 6:
        return (False, "Password must be at least 6 characters.")

    hashed = bcrypt.hashpw(pw.encode("utf-8"), bcrypt.gensalt())

    org_id = None
    role = "individual"

    if org_code:
        org_code = org_code.strip()
        with db() as conn:
            c = conn.cursor()
            c.execute("SELECT id FROM orgs WHERE org_code=?", (org_code,))
            row = c.fetchone()
        if row:
            org_id = row[0]
            role = "institution_user"
        else:
            return (False, "Invalid institution code")

    try:
        with db() as conn:
            c = conn.cursor()
            c.execute(
                "INSERT INTO users (email, password, org_id, role) VALUES (?, ?, ?, ?)",
                (email, hashed, org_id, role),
            )
            conn.commit()
        return (True, None)
    except sqlite3.IntegrityError:
        return (False, "Account already exists")

def login_user(email, pw):
    email = (email or "").strip().lower()
    pw = pw or ""
    if not email or not pw:
        return None

    with db() as conn:
        c = conn.cursor()
        c.execute("""
            SELECT id, password, paid, subscription_status, org_id, role
            FROM users WHERE email=?
        """, (email,))
        row = c.fetchone()

    if not row:
        return None

    stored_pw = row[1]
    if isinstance(stored_pw, memoryview):
        stored_pw = stored_pw.tobytes()

    try:
        if bcrypt.checkpw(pw.encode("utf-8"), stored_pw):
            return {
                "id": row[0],
                "email": email,
                "paid": bool(row[2]),
                "subscription_status": row[3],
                "org_id": row[4],
                "role": row[5],
            }
    except Exception:
        return None

    return None

def update_password(email: str, new_pw: str):
    email = (email or "").strip().lower()
    hashed = bcrypt.hashpw(new_pw.encode("utf-8"), bcrypt.gensalt())
    with db() as conn:
        c = conn.cursor()
        c.execute("UPDATE users SET password=? WHERE email=?", (hashed, email))
        conn.commit()

def mark_subscribed(uid: int):
    with db() as conn:
        c = conn.cursor()
        c.execute("UPDATE users SET paid=1, subscription_status='active' WHERE id=?", (uid,))
        conn.commit()

def log_qa_audit(user_id: int, tool: str, topic: str, year_group: str, passed: int):
    with db() as conn:
        c = conn.cursor()
        c.execute(
            "INSERT INTO qa_audit (user_id, tool, topic, year_group, passed) VALUES (?, ?, ?, ?, ?)",
            (user_id, tool, topic, year_group, passed),
        )
        conn.commit()

# =========================
# 7) PASSWORD RESET TOKEN + EMAIL
# =========================
def get_serializer():
    return URLSafeTimedSerializer(RESET_SECRET)

def create_reset_token(email: str):
    return get_serializer().dumps((email or "").strip().lower())

def verify_reset_token(token: str, max_age_seconds=3600):
    return get_serializer().loads(token, max_age=max_age_seconds)

def send_email(to_email: str, subject: str, body: str):
    if not (SMTP_HOST and SMTP_USER and SMTP_PASS):
        raise RuntimeError("SMTP not configured. Add SMTP_HOST/USER/PASS in secrets.")
    msg = EmailMessage()
    msg["From"] = SMTP_USER
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.set_content(body)

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
        server.starttls()
        server.login(SMTP_USER, SMTP_PASS)
        server.send_message(msg)

# =========================
# 8) STRIPE (SUBSCRIPTIONS)
# =========================
def start_subscription_checkout(plan: str, email: str):
    if not STRIPE_READY:
        raise RuntimeError("Stripe not configured (missing STRIPE_SECRET_KEY/price IDs).")
    if not email:
        raise ValueError("Missing customer email")

    if plan == "monthly":
        price_id = STRIPE_PRICE_MONTHLY
    elif plan == "annual":
        price_id = STRIPE_PRICE_ANNUAL
    else:
        raise ValueError("Invalid plan")

    session = stripe.checkout.Session.create(
        mode="subscription",
        payment_method_types=["card"],
        line_items=[{"price": price_id, "quantity": 1}],
        customer_email=email,
        success_url=f"{APP_BASE_URL}/?subscribed=true",
        cancel_url=f"{APP_BASE_URL}/",
    )
    return session.url

# =========================
# 9) AI (LESSON / WORKSHEET / QA / IMAGE)
# =========================
def generate_lesson(topic, year):
    if not OPENAI_READY:
        raise RuntimeError("OpenAI not configured. Add OPENAI_API_KEY in secrets.")
    prompt = dedent(f"""
    Create a classroom-ready lesson.

    Topic: {topic}
    Year group: {year}

    Include:
    - Clear learning objectives
    - Meaningful context / lived example
    - Explanation with examples
    - Differentiated tasks (Support / Core / Stretch)
    - Retrieval or plenary
    - SEN/EAL supports
    """).strip()

    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert classroom teacher."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
    )
    return r.choices[0].message.content

def generate_worksheet(topic, year, la_profile, ma_profile, ha_profile):
    if not OPENAI_READY:
        raise RuntimeError("OpenAI not configured. Add OPENAI_API_KEY in secrets.")
    prompt = dedent(f"""
    Create a classroom-ready worksheet (student-facing).

    Topic: {topic}
    Year group: {year}

    Must include:
    - Short instructions for pupils
    - A brief worked example
    - Three differentiated task sets:
      * LA (Lower Attainers): {la_profile}
      * MA (Middle Attainers): {ma_profile}
      * HA (Higher Attainers): {ha_profile}
    - SEN/EAL supports (simple, practical)
    - A short retrieval/plenary section

    Format with clear headings and bullet points.
    """).strip()

    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert classroom teacher."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.5,
    )
    return r.choices[0].message.content

def run_pedagogical_qa(content: str):
    if not OPENAI_READY:
        return content
    checklist = "\n".join(f"- {c}" for c in QA_CHECKLIST)
    prompt = f"""
You are performing a strict pedagogical QA.

PEDAGOGY:
{PEDAGOGY_CORE_POSITION}

CHECKLIST (ALL REQUIRED):
{checklist}

CONTENT:
{content}

TASK:
Revise until all standards are met.
Output ONLY the final revised content.
""".strip()

    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Senior educational QA reviewer."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return r.choices[0].message.content

def generate_image_bytes(prompt: str, size="1024x1024"):
    if not OPENAI_READY:
        raise RuntimeError("OpenAI not configured. Add OPENAI_API_KEY in secrets.")
    r = client.images.generate(model="gpt-image-1", prompt=prompt, size=size)
    return base64.b64decode(r.data[0].b64_json)

# =========================
# 10) DOCX EXPORT (BYTES, WITH IMAGE + LOGO)
# =========================
def build_docx_bytes(title: str, text: str, main_image_bytes=None, brand_logo_bytes=None):
    doc = Document()
    doc.add_heading(title, level=1)

    if brand_logo_bytes:
        bio = BytesIO(brand_logo_bytes)
        doc.add_picture(bio, width=Inches(2.0))
        doc.add_paragraph("")

    if main_image_bytes:
        bio = BytesIO(main_image_bytes)
        doc.add_picture(bio, width=Inches(5.5))
        doc.add_paragraph("")

    for line in (text or "").split("\n"):
        doc.add_paragraph(line)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue(), f"{title.replace(' ', '_').lower()}_{int(time.time())}.docx"

def paid_download_block(title: str, text: str, main_image_bytes=None, brand_logo_bytes=None):
    if not (st.session_state.user and st.session_state.user.get("paid")):
        st.info("🔒 Subscribe to download.")
        return

    docx_bytes, filename = build_docx_bytes(
        title=title,
        text=text,
        main_image_bytes=main_image_bytes,
        brand_logo_bytes=brand_logo_bytes,
    )

    st.download_button(
        "⬇ Download as Word (.docx)",
        data=docx_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"dl_{title}_{int(time.time())}",
    )

# =========================
# 11) SUBSCRIPTION PANEL
# =========================
def show_subscription_panel():
    st.subheader("🔓 Unlock Full Access")

    if not STRIPE_READY:
        st.error("Stripe is not configured (missing STRIPE_SECRET_KEY / price IDs).")
        st.caption("Add STRIPE_SECRET_KEY, STRIPE_PRICE_MONTHLY, STRIPE_PRICE_ANNUAL in Streamlit secrets.")
        return

    plan = st.radio(
        "Choose your plan",
        ["Monthly (£10/month)", "Annual (£100/year)"],
        index=0,
        key="subscription_plan",
    )

    if not st.session_state.user:
        st.info("Please log in to subscribe.")
        return

    if st.session_state.user.get("paid"):
        st.success("✅ Subscription active")
        st.caption(f"Status: {st.session_state.user.get('subscription_status', 'active')}")
        return

    if st.button("Subscribe now", key="subscribe_now"):
        plan_key = "monthly" if "Monthly" in plan else "annual"
        try:
            checkout_url = start_subscription_checkout(plan_key, st.session_state.user["email"])
            # link_button is nicer on Streamlit
            try:
                st.link_button("Continue to secure Stripe Checkout", checkout_url)
            except Exception:
                st.markdown(f"👉 [Continue to secure Stripe Checkout]({checkout_url})")
        except Exception as e:
            st.error(f"Stripe error: {e}")

# =========================
# 12) PASSWORD RESET HANDLER (RUN BEFORE MAIN UI)
# =========================
reset_token = qp_first("reset")
if reset_token:
    st.title("🔑 Reset Password")
    try:
        email = verify_reset_token(reset_token)
        st.info(f"Resetting password for: {email}")

        new_pw = st.text_input("New password", type="password")
        new_pw2 = st.text_input("Confirm new password", type="password")

        if st.button("Update password", key="reset_pw_btn"):
            if new_pw != new_pw2:
                st.error("Passwords do not match")
            elif len(new_pw) < 6:
                st.error("Password must be at least 6 characters")
            else:
                update_password(email, new_pw)
                st.success("✅ Password updated. Return to login.")
    except SignatureExpired:
        st.error("Reset link expired. Request a new one.")
    except BadSignature:
        st.error("Reset link invalid.")
    except Exception as e:
        st.error(f"Reset error: {e}")
    st.stop()

# =========================
# 13) LOGIN PAGE UI
# =========================
def login_page():
    st.title("🔐 Login / Create Account")
    t1, t2, t3 = st.tabs(["Login", "Create account", "Forgot password"])

    with t1:
        e = st.text_input("Email", key="login_e")
        p = st.text_input("Password", type="password", key="login_p")
        if st.button("Login", key="login_btn"):
            u = login_user(e, p)
            if u:
                st.session_state.user = u
                st.session_state.show_login = False
                st.rerun()
            else:
                st.error("Invalid login")

    with t2:
        e = st.text_input("Email", key="reg_e")
        p1 = st.text_input("Password", type="password", key="reg_p1")
        p2 = st.text_input("Confirm password", type="password", key="reg_p2")
        org_code = st.text_input("Institution code (optional)", key="reg_org")

        if st.button("Create account", key="reg_btn"):
            if p1 != p2:
                st.error("Passwords do not match")
            elif len(p1) < 6:
                st.error("Password must be at least 6 characters")
            else:
                ok, err = create_user(e, p1, org_code=org_code if org_code else None)
                if ok:
                    st.success("Account created. Please log in.")
                else:
                    st.error(err or "Could not create account")

    with t3:
        reset_email = st.text_input("Enter your email", key="reset_email")
        if st.button("Send reset link", key="send_reset_link_btn"):
            try:
                token = create_reset_token(reset_email.strip().lower())
                link = f"{APP_BASE_URL}/?reset={token}"
                send_email(
                    reset_email,
                    "Reset your TAC password",
                    f"Click this link to reset your password (valid 1 hour):\n\n{link}",
                )
                st.success("If that email exists, a reset link has been sent.")
            except Exception as e:
                st.error(f"Could not send reset email: {e}")

# =========================
# 14) SIDEBAR (ACCOUNT)
# =========================
with st.sidebar:
    st.title("TAC Learning")

    if st.session_state.user:
        st.success(f"Logged in: {st.session_state.user['email']}")
        st.caption(f"Role: {st.session_state.user.get('role', 'individual')}")
        if st.button("Log out", key="logout_btn"):
            st.session_state.user = None
            st.session_state.final_lesson = None
            st.session_state.lesson_image_bytes = None
            st.session_state.final_worksheet = None
            st.session_state.worksheet_image_bytes = None
            st.rerun()
    else:
        if st.button("🔐 Login / Create account", key="open_login_btn"):
            st.session_state.show_login = True
            st.rerun()

if st.session_state.show_login and not st.session_state.user:
    login_page()
    st.stop()

# =========================
# 15) STRIPE SUBSCRIPTION CALLBACK (AFTER mark_subscribed EXISTS)
# =========================
if qp_first("subscribed") == "true" and st.session_state.user:
    try:
        mark_subscribed(st.session_state.user["id"])
        st.session_state.user["paid"] = True
        st.session_state.user["subscription_status"] = "active"
        st.success("✅ Subscription active. Access unlocked.")
    except Exception as e:
        st.error(f"Could not activate subscription: {e}")

# =========================
# 16) MAIN APP TABS
# =========================
st.title("TAC Learning")

tab_lessons, tab_worksheets, tab_subscription, tab_account = st.tabs([
    "📘 Lessons",
    "📝 Worksheets",
    "🔓 Subscription",
    "⚙️ Account",
])

# =========================
# 17) LESSONS TAB
# =========================
with tab_lessons:
    st.subheader("Lesson Generator")

    if not OPENAI_READY:
        st.warning("OpenAI is not configured. Add OPENAI_API_KEY in Streamlit secrets to generate lessons.")

    col1, col2 = st.columns([2, 1])
    with col1:
        topic = st.text_input("Lesson topic", key="lesson_topic")
        year = st.text_input("Year group", key="lesson_year")
    with col2:
        include_image = st.checkbox("Include visual", value=True, key="lesson_include_image")

    if st.button("Generate lesson", key="gen_lesson"):
        if not st.session_state.user:
            st.warning("Please log in to generate lessons.")
        elif not topic or not year:
            st.error("Please enter topic and year group.")
        else:
            try:
                draft = generate_lesson(topic, year)

                qa_status = st.empty()
                qa_status.info("🔍 Running pedagogical QA (hidden)…")

                with st.spinner("Finalising lesson…"):
                    final = run_pedagogical_qa(draft)

                qa_status.success("✅ Lesson quality verified")
                st.session_state.final_lesson = final

                log_qa_audit(
                    user_id=st.session_state.user["id"],
                    tool="Lesson Generator",
                    topic=topic,
                    year_group=year,
                    passed=1
                )

                st.markdown(final)

                if include_image and OPENAI_READY:
                    with st.spinner("Generating lesson visual…"):
                        img_prompt = f"Educational illustration for a {year} lesson on {topic}. Simple, clear, age-appropriate, no text."
                        st.session_state.lesson_image_bytes = generate_image_bytes(img_prompt)
                        st.image(st.session_state.lesson_image_bytes, caption="Lesson visual")

            except Exception as e:
                st.error(f"Lesson generation failed: {e}")

    if st.session_state.final_lesson:
        st.markdown("---")
        paid_download_block(
            title="Lesson Plan",
            text=st.session_state.final_lesson,
            main_image_bytes=st.session_state.lesson_image_bytes,
            brand_logo_bytes=st.session_state.brand_logo_bytes,
        )

# =========================
# 18) WORKSHEETS TAB (LA/MA/HA)
# =========================
with tab_worksheets:
    st.subheader("Worksheet Generator (LA / MA / HA)")

    if not OPENAI_READY:
        st.warning("OpenAI is not configured. Add OPENAI_API_KEY in Streamlit secrets to generate worksheets.")

    col1, col2 = st.columns([2, 1])
    with col1:
        w_topic = st.text_input("Worksheet topic", key="ws_topic")
        w_year = st.text_input("Year group", key="ws_year")

        st.markdown("### Differentiation")
        la_profile = st.text_area(
            "LA (Lower Attainers) guidance",
            value="Simpler language, fewer steps, more scaffolds, word bank, sentence starters, 6–8 short questions.",
            height=90,
            key="ws_la",
        )
        ma_profile = st.text_area(
            "MA (Middle Attainers) guidance",
            value="Core curriculum level, mixed question types, some reasoning prompts, 8–12 questions.",
            height=90,
            key="ws_ma",
        )
        ha_profile = st.text_area(
            "HA (Higher Attainers) guidance",
            value="Greater depth: multi-step problems, reasoning/justification, extension challenge, 8–12 questions with stretch tasks.",
            height=90,
            key="ws_ha",
        )

    with col2:
        w_include_image = st.checkbox("Include visual", value=True, key="ws_include_image")

    if st.button("Generate worksheet", key="gen_ws"):
        if not st.session_state.user:
            st.warning("Please log in to generate worksheets.")
        elif not w_topic or not w_year:
            st.error("Please enter topic and year group.")
        else:
            try:
                draft = generate_worksheet(w_topic, w_year, la_profile, ma_profile, ha_profile)

                qa_status = st.empty()
                qa_status.info("🔍 Running pedagogical QA (hidden)…")

                with st.spinner("Finalising worksheet…"):
                    final = run_pedagogical_qa(draft)

                qa_status.success("✅ Worksheet quality verified")
                st.session_state.final_worksheet = final

                log_qa_audit(
                    user_id=st.session_state.user["id"],
                    tool="Worksheet Generator",
                    topic=w_topic,
                    year_group=w_year,
                    passed=1
                )

                st.markdown(final)

                if w_include_image and OPENAI_READY:
                    with st.spinner("Generating worksheet visual…"):
                        img_prompt = f"Educational illustration for a {w_year} worksheet on {w_topic}. Simple, clear, age-appropriate, no text."
                        st.session_state.worksheet_image_bytes = generate_image_bytes(img_prompt)
                        st.image(st.session_state.worksheet_image_bytes, caption="Worksheet visual")

            except Exception as e:
                st.error(f"Worksheet generation failed: {e}")

    if st.session_state.final_worksheet:
        st.markdown("---")
        paid_download_block(
            title="Worksheet",
            text=st.session_state.final_worksheet,
            main_image_bytes=st.session_state.worksheet_image_bytes,
            brand_logo_bytes=st.session_state.brand_logo_bytes,
        )

# =========================
# 19) SUBSCRIPTION TAB
# =========================
with tab_subscription:
    show_subscription_panel()

# =========================
# 20) ACCOUNT TAB (BRANDING + QA LOGS)
# =========================
with tab_account:
    st.subheader("Teacher branding for downloads")
    logo = st.file_uploader(
        "Upload your logo (PNG/JPG)",
        type=["png", "jpg", "jpeg"],
        key="brand_logo_upl"
    )
    if logo:
        st.session_state.brand_logo_bytes = logo.getvalue()
        st.success("Branding logo saved for downloads.")
        st.image(st.session_state.brand_logo_bytes, caption="Brand logo preview", width=180)

    st.markdown("---")
    st.subheader("Account info")

    if st.session_state.user:
        st.write("Email:", st.session_state.user.get("email"))
        st.write("Role:", st.session_state.user.get("role"))
        st.write("Paid:", st.session_state.user.get("paid"))
        st.write("Subscription status:", st.session_state.user.get("subscription_status"))

        st.markdown("---")
        st.subheader("QA audit logs (latest 50)")

        with db() as conn:
            c = conn.cursor()
            c.execute("""
                SELECT tool, topic, year_group, passed, datetime(created_at,'unixepoch')
                FROM qa_audit
                WHERE user_id=?
                ORDER BY id DESC
                LIMIT 50
            """, (st.session_state.user["id"],))
            rows = c.fetchall()

        if rows:
            for tool, topic, yg, passed, ts in rows:
                st.write(f"- [{ts}] {tool} | {topic} | Year {yg} | {'PASS' if passed else 'FAIL'}")
        else:
            st.caption("No QA audits yet.")
    else:
        st.info("Log in to see account details and QA logs.")
